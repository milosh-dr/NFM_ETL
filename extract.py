# docx2 ogarnij kolejność kolumn, wyrzuć jako słownik???


# Zachowaj workflow selenium dla tego systemu w oddzielnym pliku na przyszłość!


import re, os
from pathlib import Path
from io import BytesIO

import docx

from selenium import webdriver
from selenium.webdriver.common.alert import Alert
import requests
from bs4 import BeautifulSoup

from openpyxl import load_workbook

import pandas as pd
import numpy as np

# Set the home url
url = r'http://www.nfm.home.pl/orkiestra/T%20A%20B%20L%20I%20C%20A%20%20%20O%20G%20%C5%81%20O%20S%20Z%20E%20%C5%83/'
# Set the credentials for authetication
username = 'muzyk'
password = 'Coda2019!'
# Set the path to the Geckodriver executable
driver_path = '/home/milosh-dr/chromedriver'
browser_path = '/snap/bin/chromium'

def extract_web(url, driver_path, username, password):
    """
    lkjslkjd
    """

    # Set the options for the Firefox browser
    options = webdriver.ChromeOptions()
    options.add_argument('--headless')
    options.add_argument('--disable-gpu')
    options.add_argument('--remote-debugging-port=9222')
    options.binary_location = browser_path

    # service = webdriver.chrome.service.Service(executable_path=driver_path)
    os.environ['PATH'] += driver_path
    os.environ['PATH'] += browser_path

  
    # Create a new Firefox browser instance
    driver = webdriver.Chrome(options=options)
    # Use the browser to navigate to a web page
    driver.get(url)
    html = driver.page_source
    print(html)

    # Close the browser
    driver.quit()

def get_items(url, username, password):
    response = requests.get(url, auth=(username, password))
    soup = BeautifulSoup(response.text, 'html.parser')
    return soup.find_all('a')

def get_new_path(url, item, pattern):
    if re.search(pattern, item.text):
        new_path = os.path.join(url, item['href'])
        return new_path
    else:
        return None

def get_links(url, username, password):
    """
    Returns a list of dictionaries, where each dict corresponds to a particular working week.
    """
    links = []
    # Iterate over months
    for item in get_items(url, username, password):
        # Match only directories, exlude files
        path = get_new_path(url, item, '[\w\s]+/$')
        if path:
            # Iterate over weeks
            for item_2 in get_items(path, username, password):
                path_2 = get_new_path(path, item_2, '[\w\s]+/$')
                if path_2 and not re.search('ODWO', path_2):
            # DOTĄD DZIAŁA NA PEWNO
                    week_data = {}
                    week_data['lineup'] = []
                    week_data['schedule'] = []
                    # Iterate over files
                    for item_3 in get_items(path_2, username, password):
                        lineup = get_new_path(path_2, item_3, '[Ss]k[łl]ad\s*[Oo]rkiestry')
                        schedule = get_new_path(path_2, item_3, '[Pp]lan\s*[Pp]racy')
                        if lineup:
                            week_data['lineup'].append(lineup)
                        if schedule:
                            week_data['schedule'].append(schedule)
                    if week_data['lineup'] and week_data['schedule']:
                        links.append(week_data)
             
    return links

def docx_parser(docx_file_url, username, password):
    """
    Funkcja będzie działać. Musisz dodać funkcjonalność obczajania czy coś jest IPA lub PI
    Czy dienst jest wyjazdowy
    itd.
    Wyciągnąć info o dyrygencie, solistach, utworach
    """
    response = requests.get(docx_file_url, auth=(username, password))
    
    docx_bytes = BytesIO(response.content)

    docx_content = docx.Document(docx_bytes)

    text = '\n'.join([paragraph.text for paragraph in docx_content.paragraphs])
    # Extract general information about the event
    matches = re.search(r'(\d\d\d\d)\s*(?:(.+)[Pp]rogram:?\s*(.+))?', text, re.MULTILINE | re.DOTALL)
    if len(matches.groups()) > 2:
        year = matches.group(1)
        name_and_artists = matches.group(2).rstrip().split('\n')
        # Extract event name
        name = name_and_artists[0]
        # Extract artists' names
        artists = name_and_artists[1:]
        # Extract programme
        programme = [work for work in matches.group(3).replace('*', '').split('\n') if work]
    else:
        name = None
        artists = None
        programme = None
    
    tables = docx_content.tables

    # Set a list to hold dictionaries with services' information
    services = []

    for row in tables[0].rows[1:]:
        # print('-'*30)
        services_dict = {}
        date_cell = row.cells[0].text
        service_cell = row.cells[1].text
        # Extract service's date
        match = re.search('\d\d?\.\d\d?', date_cell)
        if match:
            date = match.group(0) + year

        sessions = re.findall('sesja', service_cell.lower())
        individual = re.search('ipa', service_cell.lower())
        services_total = re.findall('\d\d?\.\d\d[^\s\w\d]\d\d\.\d\d|\d\d?\.\d\d', service_cell)
        # Extract type of service
        if individual:
            continue

        else:
            services_dict['date'] = date
            services_dict['services_reg'] = len(services_total) - len(sessions)
            services_dict['services_ses'] = len(sessions)
            services.append(services_dict)

    return name, artists, programme, services

def docx_parser2(docx_file_url, username, password):
    response = requests.get(docx_file_url, auth=(username, password))
    
    docx_bytes = BytesIO(response.content)

    docx_content = docx.Document(docx_bytes)

    text = '\n'.join([paragraph.text for paragraph in docx_content.paragraphs])
    # Extract general information about the event
    matches = re.search(r'(\d\d\d\d)\s*(?:(.+)[Pp]rogram:?\s*(.+))?', text, re.MULTILINE | re.DOTALL)
    if len(matches.groups()) > 2:
        year = matches.group(1)
        name_and_artists = matches.group(2)
        if name_and_artists:
            name_and_artists = name_and_artists.rstrip().split('\n')
            # Extract event name
            name = name_and_artists[0]
            # Extract artists' names
            artists = name_and_artists[1:]
        else:
            name = None
            artists = None
        # Extract programme
        programme = matches.group(3)
        if programme:
            programme = [work for work in matches.group(3).replace('*', '').split('\n') if work]
        else:
            programme = None
    else:
        name = None
        artists = None
        programme = None
    
    tables = docx_content.tables

    table_data = []
    for row in tables[0].rows:
        row_data = []
        for cell in row.cells:
            row_data.append(cell.text.strip())
        table_data.append(row_data)

    df = pd.DataFrame(table_data[1:], columns=table_data[0])

    lowercase_cols = [col.lower() for col in df.columns]
    
    # Get indices of appropriate columns
    conductor_ind = lowercase_cols.index('dyrygent')
    hall_ind = len(lowercase_cols) - 1 - list(reversed(lowercase_cols)).index('sala')
    df = df.iloc[:,[0, 1 , hall_ind, conductor_ind]]


    matches = df.iloc[:,1].str.extractall('(\d\d?[.,:]\d\d[^\s\w\d]\d\d[.,:]\d\d|\d\d?[.,:]\d\d)').reset_index()
    # melted = pd.melt(matches, id_vars=['level_0', 'match'], value_vars=matches.columns[2:], var_name='column')
    df = pd.merge(df, matches, left_index=True, right_on='level_0')
    df['service_type'] = df.iloc[:,1].str.extract('([Ss]esja)', expand=False).fillna('Próba')
    # Account for holidays
    df.drop(df[df.iloc[:,1].str.contains('[Ww]olne', regex=True)].index, axis=0, inplace=True)

    df = df.iloc[:, [0,6,2,7,3]].copy()

    df.columns = ['date', 'time', 'hall', 'service_type', 'conductor']
    df = df.replace('[^\w]{3,}', np.nan, regex=True).fillna(method='ffill')
    df['date'] = df['date'].str.extract('(\d\d?\.\d\d?)') + '.' + year

    return name, artists, programme, df


def xlsx_parser(xlsx_file_url, username, password):
    """
    Returns a list of dictionaries, where each dictionary is a single service.
    """
    response = requests.get(xlsx_file_url, auth=(username, password))
    xlsx_bytes = BytesIO(response.content)

    columns = ['date', 'day_of_week', 'time', 'hall', 'service_type', 'order', 'conductor']
    df = pd.read_excel(xlsx_bytes, names=columns)

    text = df.iloc[0,0]
    name = None

    matches = re.search('[Pp]rogram:?\s*(.+)\s*[Ww]ykonawcy:?\s*(.+)\s*', text, re.MULTILINE | re.DOTALL)
    if matches and len(matches.groups()) > 1:
        # Extract programme
        programme = [work for work in matches.group(1).replace('*', '').split('\n') if work]
        artists = matches.group(2).split('\n')
    else:
        artists = None
        programme = None
    df = df[(~df['service_type'].isin(['PI', 'IPA', 'DW']))&(~(df['time']=='DW'))].drop(['order', 'day_of_week'], axis=1).drop([0,1,2], axis=0)
    df['date'] = df['date'].fillna(method='ffill')
    df['conductor'] = df['conductor'].fillna(method='ffill')
    df['service_type'] = df['service_type'].str.extract('(\w+)', expand=False)

    return name, artists, programme, df#.to_dict(orient='records')


def lineup_parser(xlsx_path, username, password):
    """
    Returns a list of dictonaries holding all musicians performing at any given event
    """
    # Download the XLSX file as a bytes object
    response = requests.get(xlsx_path, auth=(username, password))
    xlsx_bytes = BytesIO(response.content)

    columns = ['no', 'first_name', 'last_name']
    df = pd.read_excel(xlsx_bytes)
    # Check number of columns and read the data accordingly
    if len(df.columns) > 3:
        df1 = df.iloc[:,:3]
        df1.columns = columns
        df2 = df.iloc[:,-3:]
        df2.columns = columns
        df = pd.concat([df1, df2], axis=0)
    else:
        df.columns = columns

    mask = df[['first_name', 'last_name']].isna().all(axis=1)
    df.loc[mask, 'instrument'] = df.loc[mask, 'no']
    df['instrument'] = df['instrument'].fillna(method='ffill').str.extract('(\w+(?:\s[^\d]+)?)\s*', expand=False)
    df = df.dropna(subset=['first_name', 'last_name'], axis=0)#.drop('no', axis=1)

    return df#.to_dict(orient='records')

if __name__ == '__main__':
    # print(len(get_links(url)))
    xlsx_path = r'http://www.nfm.home.pl/orkiestra/T%20A%20B%20L%20I%20C%20A%20%20%20O%20G%20%C5%81%20O%20S%20Z%20E%20%C5%83/102%20Luty%202023/2023.01.30-02.04%20Tognetti%2Bnagranie/Sk%C5%82ad%20orkiestry%202023.01.30-02.04%20Tognetti%2Bnagranie.xls'

    # print(lineup_parser(xlsx_path))
    # links_all = get_links(url)
    # for item in links_all:
    #     print('-'*60)
    #     for plan in [i for i in item if re.search('plan pracy', i.lower())]:
    #         print(plan)
    # print(len(links_all))
    # print([item[1] for item in links_all])

    # a, b, c, services = xlsx_parser(r'http://www.nfm.home.pl/orkiestra/T%20A%20B%20L%20I%20C%20A%20%20%20O%20G%20%C5%81%20O%20S%20Z%20E%20%C5%83/102%20Luty%202023/2023.01.30-02.04%20Tognetti%2Bnagranie/Plan%20pracy%202023.01.30-02.04%20Tognetti%2Bnagranie.xlsx')
    # for service in services:
    #     print(service)

    a,b,c, services = docx_parser2(r'http://www.nfm.home.pl/orkiestra/T%20A%20B%20L%20I%20C%20A%20%20%20O%20G%20%C5%81%20O%20S%20Z%20E%20%C5%83/010%20Pa%C5%BAdziernik%202022/2022.10.17-21%20Guerrero/Plan%20pracy%202022.10.17-24%20Guerrero.docx')
    print(services)